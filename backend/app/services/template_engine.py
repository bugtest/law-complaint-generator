import json
from pathlib import Path
from typing import Dict, Any, List, Optional
from docx import Document
import io

class TemplateEngine:
    """Template management and placeholder filling engine"""

    @staticmethod
    def fill_template(
        template_path: str,
        elements: Dict[str, Any],
        output_path: str
    ) -> Dict[str, Any]:
        """
        Fill Word template with extracted elements

        Args:
            template_path: Path to the Word template
            elements: Extracted elements data
            output_path: Path to save the filled document

        Returns:
            {
                "success": bool,
                "output_path": str,
                "unmatched_placeholders": List[str],
                "error": Optional[str]
            }
        """
        try:
            doc = Document(template_path)
            unmatched = []

            # Build replacement map
            replacements = TemplateEngine._build_replacement_map(elements)

            # Replace in paragraphs
            for para in doc.paragraphs:
                TemplateEngine._replace_in_paragraph(para, replacements, unmatched)

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            TemplateEngine._replace_in_paragraph(para, replacements, unmatched)

            # Save the document
            doc.save(output_path)

            return {
                "success": True,
                "output_path": output_path,
                "unmatched_placeholders": list(set(unmatched)),
                "error": None
            }

        except Exception as e:
            return {
                "success": False,
                "output_path": "",
                "unmatched_placeholders": [],
                "error": f"模板填充失败：{str(e)}"
            }

    @staticmethod
    def _build_replacement_map(elements: Dict[str, Any]) -> Dict[str, str]:
        """Build a map of placeholder names to their values"""
        replacements = {}

        # Plaintiff
        if elements.get("plaintiff"):
            p = elements["plaintiff"]
            replacements["原告姓名"] = p.get("name", "")
            replacements["原告身份证号"] = p.get("id_number", "")
            replacements["原告地址"] = p.get("address", "")
            replacements["原告电话"] = p.get("phone", "")

        # Defendant
        if elements.get("defendant"):
            d = elements["defendant"]
            replacements["被告姓名"] = d.get("name", "")
            replacements["被告身份证号"] = d.get("id_number", "")
            replacements["被告地址"] = d.get("address", "")
            replacements["被告电话"] = d.get("phone", "")

        # Claims
        if elements.get("claims"):
            claims_text = "\n".join([
                f"{i+1}. {c.get('content', '')}"
                for i, c in enumerate(elements["claims"])
            ])
            replacements["诉讼请求"] = claims_text

        # Facts and Reasons
        replacements["事实与理由"] = elements.get("facts_and_reasons", "")

        # Evidence List
        if elements.get("evidence_list"):
            evidence_text = "\n".join([
                f"{e.get('name', '')}: {e.get('purpose', '')}"
                for e in elements["evidence_list"]
            ])
            replacements["证据清单"] = evidence_text

        return replacements

    @staticmethod
    def _replace_in_paragraph(para, replacements: Dict[str, str], unmatched: List[str]):
        """Replace placeholders in a paragraph"""
        if "{{" not in para.text:
            return

        # Find all placeholders in the paragraph
        import re
        placeholders = re.findall(r'\{\{([^}]+)\}\}', para.text)

        for placeholder in placeholders:
            placeholder_full = "{{" + placeholder + "}}"
            if placeholder in replacements:
                # Simple text replacement
                if replacements[placeholder]:
                    para.text = para.text.replace(placeholder_full, replacements[placeholder])
                else:
                    para.text = para.text.replace(placeholder_full, "[待填写]")
            else:
                unmatched.append(placeholder)
