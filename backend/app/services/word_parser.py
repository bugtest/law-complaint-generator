from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.oxml.ns import qn

class WordParser:
    """Word document parser with text extraction capabilities"""

    @staticmethod
    def extract_text(file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word document

        Returns:
            {
                "success": bool,
                "text": str,
                "paragraphs": List[str],
                "has_styles": bool,
                "error": Optional[str]
            }
        """
        try:
            doc = Document(file_path)
            paragraphs = []
            has_styles = False

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
                    # Check if paragraph has special styling
                    if para.style and para.style.name != "Normal":
                        has_styles = True

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            result = {
                "success": True,
                "text": "\n\n".join(paragraphs),
                "paragraphs": paragraphs,
                "has_styles": has_styles,
                "error": None
            }

            return result

        except Exception as e:
            return {
                "success": False,
                "text": "",
                "paragraphs": [],
                "has_styles": False,
                "error": f"Word 解析失败：{str(e)}"
            }

    @staticmethod
    def extract_placeholders(file_path: str, placeholder_pattern: str = "{{") -> Dict[str, Any]:
        """
        Extract placeholders from Word template

        Returns:
            {
                "success": bool,
                "placeholders": List[str],
                "error": Optional[str]
            }
        """
        try:
            text_result = WordParser.extract_text(file_path)
            if not text_result["success"]:
                return {"success": False, "placeholders": [], "error": text_result["error"]}

            import re
            # Find all {{xxx}} patterns
            pattern = r"\{\{([^}]+)\}\}"
            matches = re.findall(pattern, text_result["text"])

            # Clean up placeholder names
            placeholders = list(set([m.strip() for m in matches]))

            return {
                "success": True,
                "placeholders": placeholders,
                "error": None
            }

        except Exception as e:
            return {
                "success": False,
                "placeholders": [],
                "error": f"占位符提取失败：{str(e)}"
            }
