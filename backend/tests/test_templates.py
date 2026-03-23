import pytest
from fastapi.testclient import TestClient
from app.main import app
from app.database import engine, Base
from io import BytesIO
from docx import Document

@pytest.fixture(scope="function")
def test_db():
    Base.metadata.create_all(bind=engine)
    yield
    Base.metadata.drop_all(bind=engine)

@pytest.fixture
def client(test_db):
    return TestClient(app)

@pytest.fixture
def auth_token(client):
    client.post("/api/auth/register", json={
        "username": "testuser",
        "email": "test@example.com",
        "password": "password123"
    })
    login_response = client.post("/api/auth/login", json={
        "username": "testuser",
        "password": "password123"
    })
    return login_response.json()["access_token"]

def create_template_with_placeholders() -> bytes:
    bio = BytesIO()
    doc = Document()
    doc.add_paragraph("原告：{{原告姓名}}")
    doc.add_paragraph("被告：{{被告姓名}}")
    doc.add_paragraph("请求：{{诉讼请求}}")
    doc.save(bio)
    return bio.getvalue()

def test_upload_template(client, auth_token):
    file_content = create_template_with_placeholders()
    files = {"file": ("template.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}

    response = client.post(
        "/api/templates",
        files=files,
        data={"name": "民事起诉状模板"},
        headers={"Authorization": f"Bearer {auth_token}"}
    )
    assert response.status_code == 200
    data = response.json()
    assert "原告姓名" in data["placeholders"]

def test_get_templates(client, auth_token):
    response = client.get("/api/templates", headers={
        "Authorization": f"Bearer {auth_token}"
    })
    assert response.status_code == 200

def test_set_default_template(client, auth_token):
    file_content = create_template_with_placeholders()
    files = {"file": ("template.docx", file_content)}
    create_response = client.post("/api/templates", files=files, headers={"Authorization": f"Bearer {auth_token}"})
    template_id = create_response.json()["id"]

    response = client.put(f"/api/templates/{template_id}/default", headers={
        "Authorization": f"Bearer {auth_token}"
    })
    assert response.status_code == 200
    assert response.json()["is_default"] is True
