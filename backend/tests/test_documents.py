import pytest
from fastapi.testclient import TestClient
from app.main import app
from app.database import engine, Base
from io import BytesIO
from docx import Document

@pytest.fixture(scope="function")
def test_db():
    Base.metadata.create_all(bind=engine)
    yield
    Base.metadata.drop_all(bind=engine)

@pytest.fixture
def client(test_db):
    return TestClient(app)

@pytest.fixture
def auth_token(client):
    client.post("/api/auth/register", json={
        "username": "testuser",
        "email": "test@example.com",
        "password": "password123"
    })
    login_response = client.post("/api/auth/login", json={
        "username": "testuser",
        "password": "password123"
    })
    return login_response.json()["access_token"]

@pytest.fixture
def case_id(client, auth_token):
    response = client.post("/api/cases", json={
        "case_name": "Test Case"
    }, headers={"Authorization": f"Bearer {auth_token}"})
    return response.json()["id"]

def create_test_docx_content() -> bytes:
    bio = BytesIO()
    doc = Document()
    doc.add_paragraph("原告张三")
    doc.add_paragraph("被告李四")
    doc.save(bio)
    return bio.getvalue()

def test_upload_word_document(client, auth_token, case_id):
    file_content = create_test_docx_content()
    files = {"file": ("test.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}

    response = client.post(
        f"/api/cases/{case_id}/documents",
        files=files,
        data={"doc_type": "organized_word"},
        headers={"Authorization": f"Bearer {auth_token}"}
    )
    assert response.status_code == 200
    data = response.json()
    assert "id" in data

def test_get_documents(client, auth_token, case_id):
    response = client.get(f"/api/cases/{case_id}/documents", headers={
        "Authorization": f"Bearer {auth_token}"
    })
    assert response.status_code == 200
    assert isinstance(response.json(), list)

def test_parse_documents(client, auth_token, case_id):
    file_content = create_test_docx_content()
    files = {"file": ("test.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    client.post(f"/api/cases/{case_id}/documents", files=files, data={"doc_type": "organized_word"}, headers={"Authorization": f"Bearer {auth_token}"})

    response = client.post(f"/api/cases/{case_id}/parse", headers={
        "Authorization": f"Bearer {auth_token}"
    })
    assert response.status_code == 200
